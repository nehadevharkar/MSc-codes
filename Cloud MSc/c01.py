# Create a word document of your class time table and store locally and on the cloud with doc, and pdf format .


# pip install python-docx reportlab

from docx import Document

from reportlab.pdfgen import canvas


# Function to create Word document
def create_word_document(timetable, file_path):
    doc = Document()
    doc.add_heading("Class Timetable", level=1)

    # Add timetable details
    for day, schedule in timetable.items():
        doc.add_heading(day, level=2)
        for time, subject in schedule.items():
            doc.add_paragraph(f"{time}: {subject}")

    # Save the Word document
    doc.save(file_path + ".docx")


# Function to create PDF document
def create_pdf(timetable, file_path):
    c = canvas.Canvas(file_path + ".pdf")
    c.setFont("Helvetica", 12)

    # Add timetable details
    for day, schedule in timetable.items():
        c.drawString(100, c._pagesize[1] - 50, day)
        y = c._pagesize[1] - 70
        for time, subject in schedule.items():
            c.drawString(120, y, f"{time}: {subject}")
            y -= 15

    # Save the PDF document
    c.save()


# Example timetable data (replace with your actual timetable)
timetable_data = {
    "Monday": {"09:00 AM": "Math", "10:30 AM": "History"},
    "Tuesday": {"09:00 AM": "Science", "11:00 AM": "English"},
    # Add other days and subjects
}

# Specify file paths
local_file_path = "path/to/local/timetable"
cloud_file_path = "path/to/cloud/timetable"

# Create and save Word document
create_word_document(timetable_data, local_file_path)

# Create and save PDF document
create_pdf(timetable_data, local_file_path)

# Upload files to the cloud (e.g., Google Drive, Dropbox) as needed
# Your cloud storage API integration or manual upload process goes here
